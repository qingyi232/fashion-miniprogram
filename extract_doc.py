import sys
from docx import Document

doc = Document(r'F:\26dan\穿搭小程序\3 系统需求分析与总体设计.docx')

with open(r'F:\26dan\穿搭小程序\doc_output.txt', 'w', encoding='utf-8') as f:
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            f.write(f'{i}: [{para.style.name}] {t}\n')
    
    for i, table in enumerate(doc.tables):
        f.write(f'\n=== TABLE {i} ===\n')
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            f.write(' | '.join(cells) + '\n')

    f.write('\n=== IMAGES ===\n')
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            f.write(f'Image: {rel.target_ref}\n')

print('Done')
